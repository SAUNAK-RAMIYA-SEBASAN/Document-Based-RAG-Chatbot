import os
import PyPDF2
import pymupdf
from docx import Document
from typing import List, Dict, Tuple
import re
from dataclasses import dataclass


@dataclass
class DocumentChunk:
    content: str
    metadata: Dict
    chunk_id: str
    page_number: int
    section_name: str


class DocumentProcessor:
    def __init__(self, chunk_size: int = 384, overlap: int = 64):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def process_document(self, file_path: str) -> List[DocumentChunk]:
        """Process PDF or Word document into chunks"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._process_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._process_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _process_pdf(self, file_path: str) -> List[DocumentChunk]:
        """Process PDF document"""
        chunks = []
        doc = pymupdf.open(file_path)
        
        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            page_chunks = self._create_chunks(
                text, 
                os.path.basename(file_path), 
                page_num
            )
            chunks.extend(page_chunks)
        
        doc.close()
        return chunks
    
    def _process_word(self, file_path: str) -> List[DocumentChunk]:
        """Enhanced Word document processing with structure awareness"""
        chunks = []
        
        try:  
            doc = Document(file_path)
            filename = os.path.basename(file_path)
            
            # Track document structure
            current_section = "Document_Start"
            section_counter = 1
            page_counter = 1
            
            for para_idx, para in enumerate(doc.paragraphs):
                if not para.text.strip():
                    continue
                
                
                is_heading = (
                    len(para.text) < 100 and 
                    para.text.isupper() or 
                    any(style in str(para.style.name).lower() for style in ['heading', 'title'])
                )
                
                if is_heading:
                    current_section = para.text.replace(" ", "_").replace(".", "")
                    section_counter += 1
                    continue
                
                # Create chunk for regular paragraph
                chunk = DocumentChunk(
                    content=para.text.strip(),
                    metadata={'filename': filename},
                    chunk_id=f"Page_{page_counter}_{current_section}_Para_{para_idx}",
                    page_number=page_counter,
                    section_name=current_section
                )
                chunks.append(chunk)
                
                # Simulate page breaks (every 20 paragraphs)
                if para_idx % 20 == 0 and para_idx > 0:
                    page_counter += 1
            
            # Process tables separately
            for table_idx, table in enumerate(doc.tables):
                table_content = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_content += row_text + "\n"
                
                if table_content.strip():
                    chunk = DocumentChunk(
                        content=f"Table {table_idx + 1}:\n{table_content}",
                        metadata={'filename': filename},
                        chunk_id=f"Page_{page_counter}_Table_{table_idx + 1}",
                        page_number=page_counter,
                        section_name=f"Table_{table_idx + 1}"
                    )
                    chunks.append(chunk)
            
            return chunks
            
        except Exception as e:  
            print(f"Error processing Word document {file_path}: {str(e)}")
            return []

    def _create_chunks(self, text: str, filename: str, page_num: int) -> List[DocumentChunk]:
        """Create semantic chunks from text"""
        if not text.strip():
            return []
        
        chunks = []
        
        # Simple sentence-based chunking
        sentences = text.split('. ')
        current_chunk = ""
        chunk_counter = 1
        
        for sentence in sentences:
            # Add sentence to current chunk
            test_chunk = current_chunk + sentence + ". "
            
            # If chunk gets too long, save it and start new one
            if len(test_chunk.split()) > self.chunk_size:
                if current_chunk.strip():
                    chunk = DocumentChunk(
                        content=current_chunk.strip(),
                        metadata={'filename': filename},
                        chunk_id=f"Page_{page_num}_Chunk_{chunk_counter}",
                        page_number=page_num,
                        section_name=f"Section_{chunk_counter}"
                    )
                    chunks.append(chunk)
                    chunk_counter += 1
                current_chunk = sentence + ". "
            else:
                current_chunk = test_chunk
        
        # Add remaining text as final chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                metadata={'filename': filename},
                chunk_id=f"Page_{page_num}_Chunk_{chunk_counter}",
                page_number=page_num,
                section_name=f"Section_{chunk_counter}"
            )
            chunks.append(chunk)
        
        return chunks


if __name__ == "__main__":
    processor = DocumentProcessor()

